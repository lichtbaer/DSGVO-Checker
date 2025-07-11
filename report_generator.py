import json
from datetime import datetime
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from io import BytesIO
import time

class ReportGenerator:
    """
    Generates compliance reports in various formats with progress tracking
    """
    
    def __init__(self):
        """Initialize the report generator"""
        pass
    
    def generate_word_report(self, report_data: Dict[str, Any], progress_callback=None) -> bytes:
        """
        Generate a Word document report with progress tracking
        
        Args:
            report_data: Dictionary containing report data
            progress_callback: Optional callback for progress updates
            
        Returns:
            bytes: Word document as bytes
        """
        if progress_callback:
            progress_callback(0.1, "Initializing Word document...")
        
        doc = Document()
        
        # Title
        if progress_callback:
            progress_callback(0.15, "Adding document title...")
        
        title = doc.add_heading('DSGVO Compliance Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Executive Summary
        if report_data['options']['include_summary']:
            if progress_callback:
                progress_callback(0.2, "Creating executive summary...")
            
            doc.add_heading('Executive Summary', level=1)
            
            total_files = len(report_data['results'])
            avg_score = sum(r['overall_score'] for r in report_data['results']) / total_files
            
            doc.add_paragraph(f'Total Documents Analyzed: {total_files}')
            doc.add_paragraph(f'Average Compliance Score: {avg_score:.1f}%')
            
            # Compliance distribution
            scores = [r['overall_score'] for r in report_data['results']]
            compliant = len([s for s in scores if s >= 80])
            needs_improvement = len([s for s in scores if 60 <= s < 80])
            non_compliant = len([s for s in scores if s < 60])
            
            doc.add_paragraph(f'Compliant (≥80%): {compliant}')
            doc.add_paragraph(f'Needs Improvement (60-79%): {needs_improvement}')
            doc.add_paragraph(f'Non-Compliant (<60%): {non_compliant}')
            
            doc.add_paragraph()
        
        # Detailed Results
        if report_data['options']['include_details']:
            if progress_callback:
                progress_callback(0.4, "Adding detailed findings...")
            
            doc.add_heading('Detailed Findings', level=1)
            
            total_results = len(report_data['results'])
            for i, result in enumerate(report_data['results']):
                if progress_callback:
                    progress = 0.4 + (i / total_results) * 0.4
                    progress_callback(progress, f"Processing result {i + 1}/{total_results}")
                
                doc.add_heading(f'{result["filename"]} - {result["overall_score"]:.1f}%', level=2)
                
                for section, section_result in result['section_results'].items():
                    doc.add_heading(section, level=3)
                    doc.add_paragraph(f'Score: {section_result["score"]:.1f}%')
                    
                    if section_result['issues']:
                        doc.add_heading('Issues:', level=4)
                        for issue in section_result['issues']:
                            doc.add_paragraph(f'• {issue}', style='List Bullet')
                    
                    if section_result['recommendations']:
                        doc.add_heading('Recommendations:', level=4)
                        for rec in section_result['recommendations']:
                            doc.add_paragraph(f'• {rec}', style='List Bullet')
                    
                    if section_result.get('references'):
                        doc.add_heading('References (Fundstellen):', level=4)
                        for criterion, refs in section_result['references'].items():
                            doc.add_paragraph(f'- {criterion}:', style='List Bullet')
                            for ref in refs:
                                doc.add_paragraph(f'    • {ref}', style='List Bullet')
                    
                    doc.add_paragraph()
        
        # Recommendations
        if report_data['options']['include_recommendations']:
            if progress_callback:
                progress_callback(0.85, "Adding overall recommendations...")
            
            doc.add_heading('Overall Recommendations', level=1)
            
            # Collect all recommendations
            all_recommendations = []
            for result in report_data['results']:
                for section_result in result['section_results'].values():
                    all_recommendations.extend(section_result['recommendations'])
            
            # Remove duplicates and display
            unique_recommendations = list(set(all_recommendations))
            for rec in unique_recommendations:
                doc.add_paragraph(f'• {rec}', style='List Bullet')
        
        # Save to bytes
        if progress_callback:
            progress_callback(0.95, "Finalizing Word document...")
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        if progress_callback:
            progress_callback(1.0, "Word document completed!")
        
        return output.getvalue()
    
    def generate_pdf_report(self, report_data: Dict[str, Any], progress_callback=None) -> bytes:
        """
        Generate a PDF report with progress tracking
        
        Args:
            report_data: Dictionary containing report data
            progress_callback: Optional callback for progress updates
            
        Returns:
            bytes: PDF document as bytes
        """
        if progress_callback:
            progress_callback(0.1, "Initializing PDF document...")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        # Title
        if progress_callback:
            progress_callback(0.15, "Adding document title...")
        
        story.append(Paragraph('DSGVO Compliance Report', title_style))
        story.append(Paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Executive Summary
        if report_data['options']['include_summary']:
            if progress_callback:
                progress_callback(0.2, "Creating executive summary...")
            
            story.append(Paragraph('Executive Summary', styles['Heading1']))
            
            total_files = len(report_data['results'])
            avg_score = sum(r['overall_score'] for r in report_data['results']) / total_files
            
            story.append(Paragraph(f'Total Documents Analyzed: {total_files}', styles['Normal']))
            story.append(Paragraph(f'Average Compliance Score: {avg_score:.1f}%', styles['Normal']))
            
            # Compliance distribution
            scores = [r['overall_score'] for r in report_data['results']]
            compliant = len([s for s in scores if s >= 80])
            needs_improvement = len([s for s in scores if 60 <= s < 80])
            non_compliant = len([s for s in scores if s < 60])
            
            story.append(Paragraph(f'Compliant (≥80%): {compliant}', styles['Normal']))
            story.append(Paragraph(f'Needs Improvement (60-79%): {needs_improvement}', styles['Normal']))
            story.append(Paragraph(f'Non-Compliant (<60%): {non_compliant}', styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Detailed Results
        if report_data['options']['include_details']:
            if progress_callback:
                progress_callback(0.4, "Adding detailed findings...")
            
            story.append(Paragraph('Detailed Findings', styles['Heading1']))
            
            total_results = len(report_data['results'])
            for i, result in enumerate(report_data['results']):
                if progress_callback:
                    progress = 0.4 + (i / total_results) * 0.4
                    progress_callback(progress, f"Processing result {i + 1}/{total_results}")
                
                story.append(Paragraph(f'{result["filename"]} - {result["overall_score"]:.1f}%', styles['Heading2']))
                
                for section, section_result in result['section_results'].items():
                    story.append(Paragraph(section, styles['Heading3']))
                    story.append(Paragraph(f'Score: {section_result["score"]:.1f}%', styles['Normal']))
                    
                    if section_result['issues']:
                        story.append(Paragraph('Issues:', styles['Heading4']))
                        for issue in section_result['issues']:
                            story.append(Paragraph(f'• {issue}', styles['Normal']))
                    
                    if section_result['recommendations']:
                        story.append(Paragraph('Recommendations:', styles['Heading4']))
                        for rec in section_result['recommendations']:
                            story.append(Paragraph(f'• {rec}', styles['Normal']))
                    
                    if section_result.get('references'):
                        story.append(Paragraph('References (Fundstellen):', styles['Heading4']))
                        for criterion, refs in section_result['references'].items():
                            story.append(Paragraph(f'- {criterion}:', styles['Normal']))
                            for ref in refs:
                                story.append(Paragraph(f'    • {ref}', styles['Normal']))
                    
                    story.append(Spacer(1, 12))
        
        # Recommendations
        if report_data['options']['include_recommendations']:
            if progress_callback:
                progress_callback(0.85, "Adding overall recommendations...")
            
            story.append(Paragraph('Overall Recommendations', styles['Heading1']))
            
            # Collect all recommendations
            all_recommendations = []
            for result in report_data['results']:
                for section_result in result['section_results'].values():
                    all_recommendations.extend(section_result['recommendations'])
            
            # Remove duplicates and display
            unique_recommendations = list(set(all_recommendations))
            for rec in unique_recommendations:
                story.append(Paragraph(f'• {rec}', styles['Normal']))
        
        # Build PDF
        if progress_callback:
            progress_callback(0.95, "Building PDF document...")
        
        doc.build(story)
        
        if progress_callback:
            progress_callback(1.0, "PDF document completed!")
        
        buffer.seek(0)
        return buffer.getvalue() 