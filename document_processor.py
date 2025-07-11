import PyPDF2
from docx import Document
import io
import streamlit as st
from config import get_config
from utils.logger import get_logger
import time

logger = get_logger(__name__)

class DocumentProcessor:
    """
    Handles extraction of text content from various document formats
    with progress tracking for large files
    """
    
    def __init__(self):
        self.config = get_config()
    
    def extract_text(self, uploaded_file, progress_callback=None):
        """
        Extract text content from uploaded file based on its format
        
        Args:
            uploaded_file: Streamlit uploaded file object
            progress_callback: Optional callback function for progress updates
            
        Returns:
            str: Extracted text content
        """
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        try:
            if progress_callback:
                progress_callback(0.1, f"Starting extraction of {uploaded_file.name}")
            
            if file_extension == 'pdf':
                return self._extract_from_pdf(uploaded_file, progress_callback)
            elif file_extension in ['docx', 'doc']:
                return self._extract_from_docx(uploaded_file, progress_callback)
            elif file_extension == 'txt':
                return self._extract_from_txt(uploaded_file, progress_callback)
            else:
                st.error(f"Unsupported file format: {file_extension}")
                return ""
        except Exception as e:
            st.error(f"Error processing {uploaded_file.name}: {str(e)}")
            return ""
    
    def _extract_from_pdf(self, uploaded_file, progress_callback=None):
        """Extract text from PDF file with progress tracking"""
        try:
            if progress_callback:
                progress_callback(0.2, "Reading PDF file...")
            
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            total_pages = len(pdf_reader.pages)
            text_content = ""
            
            if progress_callback:
                progress_callback(0.3, f"Processing {total_pages} pages...")
            
            for page_num in range(total_pages):
                if progress_callback:
                    progress = 0.3 + (page_num / total_pages) * 0.6
                    progress_callback(progress, f"Processing page {page_num + 1}/{total_pages}")
                
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                text_content += page_text + "\n"
                
                # Add small delay for large files to prevent UI freezing
                if total_pages > 10:
                    time.sleep(0.01)
            
            if progress_callback:
                progress_callback(0.9, "Finalizing PDF extraction...")
            
            return text_content.strip()
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return ""
    
    def _extract_from_docx(self, uploaded_file, progress_callback=None):
        """Extract text from DOCX file with progress tracking"""
        try:
            if progress_callback:
                progress_callback(0.2, "Reading DOCX file...")
            
            doc = Document(uploaded_file)
            total_paragraphs = len(doc.paragraphs)
            text_content = ""
            
            if progress_callback:
                progress_callback(0.3, f"Processing {total_paragraphs} paragraphs...")
            
            for i, paragraph in enumerate(doc.paragraphs):
                if progress_callback:
                    progress = 0.3 + (i / total_paragraphs) * 0.6
                    progress_callback(progress, f"Processing paragraph {i + 1}/{total_paragraphs}")
                
                text_content += paragraph.text + "\n"
                
                # Add small delay for large files
                if total_paragraphs > 100:
                    time.sleep(0.001)
            
            if progress_callback:
                progress_callback(0.9, "Finalizing DOCX extraction...")
            
            return text_content.strip()
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    def _extract_from_txt(self, uploaded_file, progress_callback=None):
        """Extract text from TXT file with progress tracking"""
        try:
            if progress_callback:
                progress_callback(0.2, "Reading TXT file...")
            
            # Reset file pointer to beginning
            uploaded_file.seek(0)
            
            # Read file in chunks for large files
            chunk_size = 8192  # 8KB chunks
            text_content = ""
            total_size = uploaded_file.size
            bytes_read = 0
            
            while True:
                chunk = uploaded_file.read(chunk_size)
                if not chunk:
                    break
                
                text_content += chunk.decode('utf-8')
                bytes_read += len(chunk)
                
                if progress_callback:
                    progress = 0.2 + (bytes_read / total_size) * 0.7
                    progress_callback(progress, f"Reading {bytes_read}/{total_size} bytes")
            
            if progress_callback:
                progress_callback(0.9, "Finalizing TXT extraction...")
            
            return text_content.strip()
        except Exception as e:
            st.error(f"Error reading TXT: {str(e)}")
            return "" 